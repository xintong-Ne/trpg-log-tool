from __future__ import annotations

import re
from html import escape
from dataclasses import dataclass
from io import BytesIO

import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DATE_PATTERN = r"\d{4}[-/]\d{2}[-/]\d{2}"
TIME_PATTERN = r"\d{2}:\d{2}:\d{2}"
DATE_TIME_PATTERN = rf"(?:{DATE_PATTERN}(?:\s+{TIME_PATTERN})?|{TIME_PATTERN})"

NAME_FIRST_HEADER_RE = re.compile(rf"^\s*(?P<name>.+?)\s+(?P<time>{DATE_TIME_PATTERN})\s*$")
TIME_FIRST_HEADER_RE = re.compile(rf"^\s*(?P<time>{DATE_PATTERN}(?:\s+{TIME_PATTERN})?)\s+(?P<name>.+?)\s*$")
INLINE_HEADER_RE = re.compile(rf"^\s*(?P<time>{DATE_TIME_PATTERN})\s+<(?P<name>[^>]+)>\s*(?P<content>.*)$")
COLON_INLINE_HEADER_RE = re.compile(rf"^\s*(?P<name>.+?)\s+(?P<time>{DATE_TIME_PATTERN})\s*[：:]\s*(?P<content>.*)$")
DATE_ONLY_RE = re.compile(rf"^{DATE_PATTERN}$")
TIME_ONLY_RE = re.compile(rf"^{TIME_PATTERN}$")
TRAILING_QQ_RE = re.compile(r"^(?P<name>.*?)\s*[\(（](?P<qq>\d+)[\)）]\s*$")

@dataclass
class SourceInput:
    name: str
    raw: str


@dataclass
class ChatRecord:
    timestamp: tuple[int, int, int, int, int, int]
    file_index: int
    record_index: int
    filename: str
    name: str
    qq: str
    time_text: str
    content: str
    text: str


@dataclass
class HeadingItem:
    filename: str
    title: str


@dataclass
class OutputOptions:
    remove_qq: bool
    remove_time: bool
    remove_bracket_records: bool
    remove_image_records: bool
    add_file_headings: bool
    format_brackets: bool
    color_brackets: bool
    bracket_color: str


@dataclass
class SpeakerSettings:
    display_name: str
    text_color: str
    delete: bool


def clean_hex(value: str) -> str:
    return value.strip().lstrip("#").upper()


def decode_upload(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "utf-16le", "utf-16be", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def sortable_timestamp(value: str, day_offset: int = 0) -> tuple[int, int, int, int, int, int]:
    if TIME_ONLY_RE.match(value):
        hour, minute, second = (int(part) for part in value.split(":"))
        return (1970, 1, 1 + day_offset, hour, minute, second)
    parts = [int(part) for part in re.split(r"[-/ :]", value)]
    if DATE_ONLY_RE.match(value):
        year, month, day = parts
        return (year, month, day, 0, 0, 0)
    return tuple(parts)


def normalize_header_name(raw_name: str) -> tuple[str, str]:
    name = re.sub(r"\s+", " ", raw_name.strip())
    qq = ""
    qq_match = TRAILING_QQ_RE.match(name)
    if qq_match:
        name = re.sub(r"\s+", " ", qq_match.group("name").strip())
        qq = qq_match.group("qq")
    return name, qq


def parse_record_start(line: str) -> tuple[str, str, str, str | None] | None:
    for header_re in (INLINE_HEADER_RE, COLON_INLINE_HEADER_RE, TIME_FIRST_HEADER_RE, NAME_FIRST_HEADER_RE):
        match = header_re.match(line)
        if not match:
            continue
        name, qq = normalize_header_name(match.group("name"))
        return name, qq, match.group("time"), match.groupdict().get("content")
    return None


def format_record(name: str, qq: str, time_text: str, content: str, options: OutputOptions) -> str:
    header = name
    if qq and not options.remove_qq:
        header += f"({qq})"
    if not options.remove_time:
        header += f" {time_text}"
    return f"{header}：{content}"


def with_display_name(record: ChatRecord, display_name: str, options: OutputOptions) -> ChatRecord:
    return ChatRecord(
        timestamp=record.timestamp,
        file_index=record.file_index,
        record_index=record.record_index,
        filename=record.filename,
        name=record.name,
        qq=record.qq,
        time_text=record.time_text,
        content=record.content,
        text=format_record(display_name, record.qq, record.time_text, record.content, options),
    )


def should_keep_record(content: str, options: OutputOptions) -> bool:
    stripped = content.strip()
    if options.remove_image_records and (stripped == "[图片]" or stripped.startswith("[CQ:image,file=https:")):
        return False
    if options.remove_bracket_records and stripped.startswith(("（", "(")):
        return False
    return True


def read_records(raw: str, filename: str, file_index: int, options: OutputOptions) -> list[ChatRecord]:
    raw = raw.replace("\r\n", "\n").replace("\r", "\n")
    lines = raw.split("\n")
    starts = [
        idx
        for idx, line in enumerate(lines)
        if parse_record_start(line) is not None
    ]
    if not starts:
        raise ValueError(f"{filename} 没有找到符合格式的聊天记录开头。")

    records = []
    day_offset = 0
    previous_time_key: tuple[int, int, int] | None = None
    for record_index, start in enumerate(starts):
        end = starts[record_index + 1] if record_index + 1 < len(starts) else len(lines)
        block_lines = lines[start:end]
        while block_lines and not block_lines[-1].strip():
            block_lines.pop()

        start_parts = parse_record_start(block_lines[0])
        if start_parts is None:
            raise ValueError(f"无法解析记录开头：{block_lines[0]}")
        name, qq, time_text, inline_content = start_parts
        if inline_content is None:
            content_lines = block_lines[1:]
        else:
            content_lines = [inline_content, *block_lines[1:]]
        if TIME_ONLY_RE.match(time_text):
            time_key = tuple(int(part) for part in time_text.split(":"))
            if previous_time_key is not None and time_key < previous_time_key:
                day_offset += 1
            previous_time_key = time_key
        else:
            previous_time_key = None
        timestamp = sortable_timestamp(time_text, day_offset)
        content = "\n".join(content_lines).strip("\n")
        if not should_keep_record(content, options):
            continue

        records.append(
            ChatRecord(
                timestamp=timestamp,
                file_index=file_index,
                record_index=record_index,
                filename=filename,
                name=name,
                qq=qq,
                time_text=time_text,
                content=content,
                text=format_record(name, qq, time_text, content, options),
            )
        )
    return records


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), clean_hex(fill))


def is_bracket_record(record: ChatRecord) -> bool:
    return record.content.strip().startswith(("（", "("))


def add_styled_run(paragraph, text: str, color: str, size: Pt, italic: bool) -> None:
    run = paragraph.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = size
    run.font.italic = True if italic else None
    run.font.color.rgb = RGBColor.from_string(clean_hex(color))


def add_preserved_text(paragraph, text: str, font_color: str, size: Pt, italic: bool) -> None:
    parts = text.split("\n")
    for idx, part in enumerate(parts):
        if idx:
            paragraph.add_run().add_break()
        if part:
            add_styled_run(paragraph, part, font_color, size, italic)


def add_heading(document: Document, title: str) -> None:
    if not title:
        add_white_spacer(document)
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0
    set_paragraph_shading(paragraph, "#FFFFFF")
    run = paragraph.add_run(title)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("000000")


def add_white_spacer(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    set_paragraph_shading(paragraph, "#FFFFFF")
    run = paragraph.add_run(" ")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string("FFFFFF")


def build_docx(
    items: list[ChatRecord | HeadingItem],
    file_fill_colors: dict[str, str],
    speaker_settings: dict[str, SpeakerSettings],
    options: OutputOptions,
) -> BytesIO:
    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(42)
    section.right_margin = Pt(42)

    normal = document.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10.5)

    for index, item in enumerate(items):
        if isinstance(item, HeadingItem):
            add_heading(document, item.title)
            continue

        record = item
        fill_color = file_fill_colors[record.filename]
        text_color = speaker_settings.get(record.name, SpeakerSettings(record.name, "#000000", False)).text_color
        bracket_record = is_bracket_record(record)
        if bracket_record and options.color_brackets:
            text_color = options.bracket_color
        size = Pt(9) if bracket_record and options.format_brackets else Pt(10.5)
        italic = bool(bracket_record and options.format_brackets)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        set_paragraph_shading(paragraph, fill_color)
        add_preserved_text(paragraph, record.text, text_color, size, italic)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def collect_records(sources: list[SourceInput], options: OutputOptions) -> list[ChatRecord]:
    records: list[ChatRecord] = []
    for file_index, source in enumerate(sources):
        records.extend(read_records(source.raw, source.name, file_index, options))
    records.sort(key=lambda item: (item.timestamp, item.file_index, item.record_index))
    return records


def build_output_items(
    records: list[ChatRecord],
    options: OutputOptions,
    file_heading_titles: dict[str, str],
    speaker_settings: dict[str, SpeakerSettings],
) -> list[ChatRecord | HeadingItem]:
    records = [
        with_display_name(record, speaker_settings[record.name].display_name, options)
        for record in records
        if record.name in speaker_settings and not speaker_settings[record.name].delete
    ]

    if not options.add_file_headings:
        return list(records)

    items: list[ChatRecord | HeadingItem] = []
    current_filename = None
    for record in records:
        if record.filename != current_filename:
            items.append(HeadingItem(record.filename, file_heading_titles.get(record.filename, "")))
            current_filename = record.filename
        items.append(record)
    return items


def preview_records(
    items: list[ChatRecord | HeadingItem],
    file_fill_colors: dict[str, str],
    speaker_settings: dict[str, SpeakerSettings],
    options: OutputOptions,
) -> None:
    st.subheader("预览")
    if not items:
        st.info("这里会显示染色后的效果。")
        return

    preview_max = min(2000, len(items))
    if preview_max <= 20:
        preview_limit = preview_max
    else:
        preview_limit = st.slider(
            "预览数量",
            min_value=20,
            max_value=preview_max,
            value=min(200, preview_max),
            step=20,
        )
    st.caption(f"当前显示前 {preview_limit} 项，共 {len(items)} 项。生成 docx 时会处理全部内容。")

    blocks = []
    for item in items[:preview_limit]:
        if isinstance(item, HeadingItem):
            if item.title:
                blocks.append(f"<div class=\"preview-heading\">{escape(item.title)}</div>")
            else:
                blocks.append("<div class=\"preview-gap\"></div>")
            continue

        record = item
        text = escape(record.text).replace("\n", "<br>")
        fill_color = file_fill_colors.get(record.filename, "#FFFFFF")
        text_color = speaker_settings.get(record.name, SpeakerSettings(record.name, "#000000", False)).text_color
        styles = []
        if is_bracket_record(record):
            if options.color_brackets:
                text_color = options.bracket_color
            if options.format_brackets:
                styles.extend(["font-size: 0.9em", "font-style: italic"])
        style_suffix = ";" + ";".join(styles) if styles else ""
        blocks.append(
            f"""
            <div class="preview-record" style="background:{fill_color};color:{text_color}{style_suffix};">
                {text}
            </div>
            """
        )

    components.html(
        f"""
        <!doctype html>
        <html>
        <head>
        <meta charset="utf-8">
        <style>
        body {{
            margin: 0;
            font-family: SimSun, Songti SC, serif;
            background: #ffffff;
        }}
        .preview-box {{
            border: 1px solid #d9dde3;
            background: #ffffff;
            height: 800px;
            overflow-y: auto;
            padding: 12px;
            box-sizing: border-box;
        }}
        .preview-record {{
            padding: 8px 10px;
            border-radius: 4px;
            font-size: 14px;
            line-height: 1.55;
            white-space: normal;
            overflow-wrap: anywhere;
        }}
        .preview-gap {{
            height: 12px;
            background: #ffffff;
        }}
        .preview-heading {{
            background: #ffffff;
            color: #000000;
            font-weight: 700;
            text-align: center;
            padding: 7px 10px;
            margin: 6px 0;
        }}
        </style>
        </head>
        <body>
        <div class="preview-box">
            {''.join(blocks)}
        </div>
        </body>
        </html>
        """,
        height=830,
        scrolling=False,
    )


def main() -> None:
    st.set_page_config(page_title="跑团 Log 整理工具", page_icon="📄", layout="wide")
    st.markdown(
        """
        <div style="display:flex;align-items:baseline;gap:10px;margin-bottom:0.5rem;">
            <h1 style="margin:0;">跑团 Log 整理工具</h1>
            <span style="font-size:0.9rem;color:#6b7280;">by 仿生橘猫</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    input_mode = st.radio("记录来源", ["上传 txt 文件", "粘贴记录"], horizontal=True)
    sources: list[SourceInput] = []

    if input_mode == "上传 txt 文件":
        uploaded_files = st.file_uploader(
            "上传 txt log 文件",
            type=["txt"],
            accept_multiple_files=True,
        )
        for uploaded_file in uploaded_files:
            sources.append(SourceInput(uploaded_file.name, decode_upload(uploaded_file)))
    else:
        paste_count = st.number_input("粘贴来源数量", min_value=1, max_value=12, value=1, step=1)
        for index in range(int(paste_count)):
            with st.expander(f"粘贴来源 {index + 1}", expanded=index == 0):
                name = st.text_input("来源名称", value=f"粘贴记录 {index + 1}", key=f"paste-name-{index}")
                raw = st.text_area("记录内容", height=220, key=f"paste-raw-{index}")
                if raw.strip():
                    sources.append(SourceInput(name.strip() or f"粘贴记录 {index + 1}", raw))

    if not sources:
        st.info("上传 txt 文件，或者粘贴一份记录后，就可以设置颜色并生成 docx。")
        return

    settings_col, preview_col = st.columns([0.75, 1], gap="large")

    records: list[ChatRecord] = []
    items: list[ChatRecord | HeadingItem] = []
    file_fill_colors: dict[str, str] = {}
    file_heading_titles: dict[str, str] = {}
    speaker_settings: dict[str, SpeakerSettings] = {}
    with settings_col:
        st.subheader("输出选项")
        option_cols = st.columns(2)
        options = OutputOptions(
            remove_qq=option_cols[0].checkbox("去除 QQ 号", value=True),
            remove_time=option_cols[1].checkbox("去除时间", value=True),
            remove_bracket_records=option_cols[0].checkbox("去除括号开头内容", value=False),
            remove_image_records=option_cols[1].checkbox("去除 [图片]", value=False),
            add_file_headings=option_cols[0].checkbox("按文件分段加小标题", value=False),
            format_brackets=option_cols[1].checkbox("括号内容小一号 + 斜体", value=False),
            color_brackets=option_cols[0].checkbox("括号内容单独染色", value=False),
            bracket_color=option_cols[1].color_picker("括号颜色", value="#808080"),
        )

        try:
            records = collect_records(sources, options)
        except Exception as exc:
            st.error(str(exc))
            return

        st.subheader("文件底色")
        for index, source in enumerate(sources):
            cols = st.columns([3, 1])
            cols[0].markdown(f"**{source.name}**")
            fill_color = cols[1].color_picker(
                "底色",
                value="#FFFFFF",
                key=f"fill-{index}-{source.name}",
            )
            file_fill_colors[source.name] = fill_color

        if options.add_file_headings:
            st.subheader("文件分段小标题")
            st.caption("留空时会插入一个白色空行。小标题会在排序后每次切换来源时出现。")
            for index, source in enumerate(sources):
                file_heading_titles[source.name] = st.text_input(
                    f"{source.name} 的小标题",
                    value=source.name,
                    key=f"heading-{index}-{source.name}",
                )

        st.subheader("人物设置")
        speaker_names = sorted({record.name for record in records})
        if not speaker_names:
            st.info("当前选项过滤后没有可输出的记录。")
        for index, speaker_name in enumerate(speaker_names):
            cols = st.columns([2.1, 2.1, 1, 1.6])
            cols[0].markdown(f"**{speaker_name}**")
            display_name = cols[1].text_input(
                "显示名",
                value=speaker_name,
                key=f"name-{index}-{speaker_name}",
            )
            text_color = cols[2].color_picker(
                "文字",
                value="#000000",
                key=f"text-{index}-{speaker_name}",
            )
            delete = cols[3].checkbox(
                "删除",
                value=False,
                key=f"delete-{index}-{speaker_name}",
            )
            speaker_settings[speaker_name] = SpeakerSettings(display_name, text_color, delete)

        items = build_output_items(records, options, file_heading_titles, speaker_settings)

        output_name = st.text_input("输出文件名", value="合并log.docx")
        if not output_name.lower().endswith(".docx"):
            output_name = f"{output_name}.docx"

        if st.button("生成 docx", type="primary"):
            try:
                docx_file = build_docx(items, file_fill_colors, speaker_settings, options)
            except Exception as exc:
                st.error(str(exc))
                return

            st.success(f"已合并 {len(records)} 条记录。")
            st.download_button(
                "下载 docx",
                data=docx_file,
                file_name=output_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    with preview_col:
        preview_records(items, file_fill_colors, speaker_settings, options)


if __name__ == "__main__":
    main()
