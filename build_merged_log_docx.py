from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "流水线心脏_合并log.docx"

HEADER_RE = re.compile(r"^.+\(\d+\) \d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}$")
HEADER_PARTS_RE = re.compile(r"^(?P<name>.+)\((?P<qq>\d+)\) (?P<time>\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2})$")
TIME_RE = re.compile(r"(\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2})$")

FILES = [
    "g515493694_My Heart Is Gold (1).txt",
    "g937953306_Cocktail Molotov.txt",
    "g515701790_Secret in The Moonlight.txt",
    "g779188663_Take Me To Church.txt",
    "g955685548_流水线心脏.txt",
    "g955685548_流水线心脏 下.txt",
    "g955685548_流水线心脏 下下.txt",
    "g955685548_流水线心脏.temp.txt",
]

STYLES = {
    "g515493694_My Heart Is Gold (1).txt": ("FFFFFF", "75B7BA"),
    "g937953306_Cocktail Molotov.txt": ("FFFFFF", "4E8259"),
    "g515701790_Secret in The Moonlight.txt": ("FFFFFF", "BB9EC5"),
    "g779188663_Take Me To Church.txt": ("FFFFFF", "638099"),
}

EXCLUDED_COLOR_QQS = {"3645066195", "3591750778"}
WHITE_FILL_QQ_TEXT_COLORS = {
    "3645066195": "B39F8F",
    "3591750778": "949494",
}
WHITE_FILL_NAME_TEXT_COLORS = {
    "赫斯提亚雪夜": "B39F8F",
    "梅迪奇": "949494",
}
SEGMENT_START = "一块电子脑：你在白茫茫的世界中沿着唯一的那条道路奔跑着，追逐着前方的人影。"
SEGMENT_END = "一颗电子心：“诺沃蒙多和新芝加哥的边界。”昆廷定好导航，水谷隼在你的控制下驰骋而去。"


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def add_preserved_text(paragraph, text: str, font_color: str) -> None:
    parts = text.split("\n")
    for idx, part in enumerate(parts):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.font.name = "SimSun"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(font_color)


def sortable_timestamp(value: str) -> tuple[int, int, int, int, int, int]:
    return tuple(int(part) for part in re.split(r"[- :]", value))


def format_record(block_lines: list[str]) -> str:
    name = parse_header(block_lines[0])[0]
    content = "\n".join(block_lines[1:])
    return f"{name}：{content}"


def parse_header(header: str) -> tuple[str, str, str]:
    match = HEADER_PARTS_RE.match(header)
    if not match:
        raise ValueError(f"Invalid record header: {header!r}")
    return match.group("name"), match.group("qq"), match.group("time")


def read_records(path: Path, file_index: int):
    raw = path.read_text(encoding="utf-8-sig")
    raw = raw.replace("\r\n", "\n").replace("\r", "\n")
    lines = raw.split("\n")
    starts = [idx for idx, line in enumerate(lines) if HEADER_RE.match(line)]
    if not starts:
        raise ValueError(f"No records found in {path.name}")

    records = []
    for rec_index, start in enumerate(starts):
        end = starts[rec_index + 1] if rec_index + 1 < len(starts) else len(lines)
        block_lines = lines[start:end]
        while block_lines and block_lines[-1] == "":
            block_lines.pop()
        speaker_name, speaker_qq, time_text = parse_header(block_lines[0])
        timestamp = sortable_timestamp(time_text)
        block = format_record(block_lines)
        records.append((timestamp, file_index, rec_index, path.name, block, speaker_name, speaker_qq))
    return records


def build_player_text_colors() -> dict[str, str]:
    colors = {}
    for file_index, filename in enumerate(FILES[:4]):
        fill = STYLES[filename][1]
        for record in read_records(ROOT / filename, file_index):
            speaker_qq = record[6]
            if speaker_qq not in EXCLUDED_COLOR_QQS:
                colors[speaker_qq] = fill
    return colors


def group_segment_by_source(records):
    start_index = next((idx for idx, record in enumerate(records) if SEGMENT_START in record[4]), None)
    end_index = next((idx for idx, record in enumerate(records) if SEGMENT_END in record[4]), None)
    if start_index is None:
        raise ValueError("Could not find segment start record")
    if end_index is None:
        raise ValueError("Could not find segment end record")
    if start_index > end_index:
        raise ValueError("Segment start is after segment end")

    before = records[:start_index]
    segment = records[start_index : end_index + 1]
    after = records[end_index + 1 :]

    grouped = []
    seen_file_indexes = []
    for record in segment:
        file_index = record[1]
        if file_index not in seen_file_indexes:
            seen_file_indexes.append(file_index)

    start_file_index = segment[0][1]
    end_file_index = segment[-1][1]
    ordered_file_indexes = [start_file_index]
    ordered_file_indexes.extend(
        file_index
        for file_index in seen_file_indexes
        if file_index not in {start_file_index, end_file_index}
    )
    if end_file_index != start_file_index:
        ordered_file_indexes.append(end_file_index)

    for file_index in ordered_file_indexes:
        grouped.extend(record for record in segment if record[1] == file_index)

    print(f"Grouped segment records: {len(segment)}")
    print("Grouped segment sources:")
    for file_index in ordered_file_indexes:
        count = sum(1 for record in segment if record[1] == file_index)
        filename = next(record[3] for record in segment if record[1] == file_index)
        print(f"  {filename}: {count}")

    return before + grouped + after


def build() -> None:
    all_records = []
    for file_index, filename in enumerate(FILES):
        all_records.extend(read_records(ROOT / filename, file_index))

    all_records.sort(key=lambda item: (item[0], item[1], item[2]))
    all_records = group_segment_by_source(all_records)
    player_text_colors = build_player_text_colors()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(42)
    section.right_margin = Pt(42)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10.5)

    for record in all_records:
        filename = record[3]
        block = record[4]
        font_color, fill = STYLES.get(filename, ("000000", "FFFFFF"))
        if fill == "FFFFFF":
            font_color = (
                WHITE_FILL_QQ_TEXT_COLORS.get(record[6])
                or WHITE_FILL_NAME_TEXT_COLORS.get(record[5])
                or player_text_colors.get(record[6])
                or font_color
            )
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.0
        set_paragraph_shading(paragraph, fill)
        add_preserved_text(paragraph, block, font_color)

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Records: {len(all_records)}")


if __name__ == "__main__":
    build()
